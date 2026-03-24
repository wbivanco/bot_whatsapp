import os

from docx import Document as DocxReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document as LCDocument


class DocxPythonLoader:
    """Carga .docx con python-docx (sin dependencia ``unstructured``)."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        doc = DocxReader(self.file_path)
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = " | ".join((c.text or "").strip() for c in row.cells)
                if cells.strip():
                    rows.append(cells)
            if rows:
                parts.append("\n".join(rows))
        body = "\n\n".join(parts)
        return [LCDocument(page_content=body, metadata={})]


class Utils:
    @staticmethod
    def generate_chunks(documento, chunk_size=1000, chunk_overlap=200):
        """ Divide un documento en fragmentos de texto. """
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""]  # Separadores más específicos
        )
        document_chunks = text_splitter.split_documents(documento)

        return document_chunks

    @staticmethod
    def load_chunked_documents(upload_directory, file_types):
        """ Genera una lista de chunks de documentos a partir de los archivos subidos al directorio de uploads y con las 
        extensiones permitidas pasadas como parámetro. """                
        loader_mapping = {
            ".pdf": PyPDFLoader,
            ".docx": DocxPythonLoader,
            ".txt": TextLoader,
        }

        documents = []
        allowed = {ext.lower() for ext in file_types}
        for filename in os.listdir(upload_directory):
            file_path = os.path.join(upload_directory, filename)
            file_extension = os.path.splitext(filename)[1].lower()

            if file_extension in allowed:
                loader_class = loader_mapping.get(file_extension)
                if loader_class:
                    loader = loader_class(file_path)
                    documento = loader.load()

                    for doc in documento:
                        doc.metadata["source"] = filename
                    total_text = sum(len((d.page_content or "").strip()) for d in documento)
                    if total_text == 0:
                        print(
                            f"[Embeddings] Advertencia: '{filename}' no aportó texto extraíble "
                            f"(PDF escaneado/imagen u otro formato sin capa de texto). No se indexará contenido útil."
                        )
                    document_chunks = Utils.generate_chunks(documento)
                    documents.extend(document_chunks)
                else:
                    print(f"No se encontró un cargador para el tipo de archivo: {file_extension}")
            else:
                print(f"Tipo de archivo no soportado: {filename}")

        return documents
    